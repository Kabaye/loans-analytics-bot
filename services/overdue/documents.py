from __future__ import annotations

from datetime import datetime, timedelta
from pathlib import Path
import json
import logging
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from bot import config
from bot.utils.borrower_addresses import merge_primary_borrower_address, normalize_borrower_addresses

log = logging.getLogger(__name__)

SERVICE_URLS = {
    "kapusta": "https://kapusta.by",
    "finkit": "https://finkit.by",
    "zaimis": "https://zaimis.by",
}

GENERATED_DOCS_DIR = Path(config.BASE_DIR) / "data" / "generated-docs"
GENERATED_DOCS_TTL_DAYS = 5
SMS_MAX_LEN = 140
CLAIM_VOLUNTARY_TERM_DAYS = 7
CLAIM_FONT_SIZE_PT = 13
SMS_SERVICE_NAMES = {
    "finkit": "ФинКит",
    "zaimis": "ЗАЙМись",
    "kapusta": "Kapusta",
}
_UUID_RE = re.compile(
    r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$",
    re.IGNORECASE,
)
ADDRESS_ABBREVIATIONS = (
    (r"^город\s+", "г. "),
    (r"^г\.?\s+", "г. "),
    (r"^деревня\s+", "д. "),
    (r"^д\.?\s+", "д. "),
    (r"^агрогородок\s+", "аг. "),
    (r"^аг\.?\s+", "аг. "),
    (r"^пос[её]лок\s+", "пос. "),
    (r"^пос\.?\s+", "пос. "),
    (r"^улица\s+", "ул. "),
    (r"^ул\.?\s+", "ул. "),
    (r"^проспект\s+", "пр-т "),
    (r"^пр\-т\s+", "пр-т "),
    (r"^переулок\s+", "пер. "),
    (r"^пер\.?\s+", "пер. "),
    (r"^дом\s+", "д. "),
    (r"^д\.?\s+", "д. "),
    (r"^квартира\s+", "кв. "),
    (r"^кв\.?\s+", "кв. "),
)
LOCALITY_PREFIXES = ("г.", "город ", "д.", "деревня ", "аг.", "агрогородок ", "пос.", "поселок ", "посёлок ")


def _money(value: float | int | None) -> str:
    try:
        return f"{float(value or 0):.2f} BYN"
    except (TypeError, ValueError):
        return "0.00 BYN"


def _number(value) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _parse_date(value: str | None) -> datetime | None:
    if not value:
        return None
    for candidate in (value, value.replace("Z", "+00:00")):
        try:
            return datetime.fromisoformat(candidate)
        except ValueError:
            continue
    return None


def _date(value: str | None) -> str:
    dt = _parse_date(value)
    if dt:
        return dt.strftime("%d.%m.%Y")
    return value or "—"


def _service_url(case: dict) -> str:
    return SERVICE_URLS.get(case.get("service"), "")


def _coalesce_text(*values: object | None) -> str | None:
    for value in values:
        text = str(value or "").strip()
        if text:
            return text
    return None


def _zaimis_order_code(case: dict) -> str | None:
    if str(case.get("service") or "").lower() != "zaimis":
        return None
    payload = _parse_case_raw(case)
    detail = payload.get("detail") or {}
    order = payload.get("order") or payload.get("list") or {}
    offer = detail.get("offer") or order.get("offer") or {}
    code = _coalesce_text(
        detail.get("code"),
        order.get("code"),
        offer.get("code"),
    )
    return code


def _looks_like_uuid(value: str | None) -> bool:
    text = str(value or "").strip()
    return bool(text and _UUID_RE.fullmatch(text))


def _loan_ref(case: dict) -> str:
    loan_number = _coalesce_text(case.get("loan_number"))
    loan_id = _coalesce_text(case.get("loan_id"))
    external_id = _coalesce_text(case.get("external_id"))
    zaimis_code = _zaimis_order_code(case)
    if zaimis_code and (
        not loan_number
        or _looks_like_uuid(loan_number)
        or loan_number == loan_id
        or loan_number == external_id
    ):
        return zaimis_code
    return loan_number or loan_id or external_id or "—"


def build_case_loan_ref(case: dict) -> str:
    return _loan_ref(case)


def _claim_deadline_date() -> str:
    return (datetime.now() + timedelta(days=CLAIM_VOLUNTARY_TERM_DAYS)).strftime("%d.%m.%Y")


def _sms_name(full_name: str | None) -> str:
    parts = [part for part in str(full_name or "").strip().split() if part]
    if not parts:
        return "Должник"
    if len(parts) == 1:
        return parts[0][:20]
    surname = parts[0].title()[:18]
    initials = "".join(part[0].upper() for part in parts[1:] if part)
    return f"{surname} {initials}".strip()[:24]


def _sms_ref(case: dict) -> str:
    ref = _loan_ref(case)
    return ref if len(ref) <= 8 else ref[-5:]


def _sms_date(case: dict) -> str:
    dt = _parse_date(case.get("issued_at"))
    if dt:
        return dt.strftime("%d.%m.%y")
    text = case.get("issued_at") or ""
    return text[:10] if text else "—"


def _fit_sms(text: str) -> str:
    compact = " ".join(text.split()).strip()
    if len(compact) <= SMS_MAX_LEN:
        return compact
    trimmed = compact[:SMS_MAX_LEN + 1]
    cut_at = trimmed.rfind(" ")
    if cut_at >= 90:
        return trimmed[:cut_at].rstrip(" ,.-")
    return compact[:SMS_MAX_LEN].rstrip(" ,.-")


def _sms_service(case: dict) -> str:
    return SMS_SERVICE_NAMES.get(str(case.get("service") or "").lower(), str(case.get("service") or "займ"))


def _sms_amount(case: dict) -> str:
    try:
        value = float(case.get("total_due") or 0)
    except (TypeError, ValueError):
        return "0р"
    if value.is_integer():
        return f"{int(value)}р"
    text = f"{value:.2f}".rstrip("0").rstrip(".")
    return f"{text}р"


def _address_line(zip_code: str | None, address: str | None) -> str:
    parts = []
    if zip_code:
        parts.append(str(zip_code).strip())
    if address:
        parts.append(str(address).strip())
    line = ", ".join(part for part in parts if part)
    if "беларус" not in line.lower():
        line = f"Республика Беларусь, {line}" if line else "Республика Беларусь"
    return line


def _parse_case_raw(case: dict) -> dict:
    raw = case.get("raw_data")
    if isinstance(raw, dict):
        return raw
    if not raw:
        return {}
    try:
        return json.loads(raw)
    except Exception:
        return {}


def _case_borrower_addresses(case: dict) -> list[dict[str, str]]:
    payload = _parse_case_raw(case)
    overrides = payload.get("contact_overrides") or {}
    raw_addresses = normalize_borrower_addresses((payload.get("expired_info") or {}).get("addressStr"))
    return merge_primary_borrower_address(
        case.get("borrower_address"),
        case.get("borrower_zip"),
        case.get("borrower_addresses") or overrides.get("borrower_addresses") or raw_addresses,
        full_name=case.get("full_name"),
    )


def _format_address_with_zip(address_entry: dict[str, str]) -> str:
    address = str(address_entry.get("address") or "").strip()
    zip_code = str(address_entry.get("zip") or "").strip()
    if not address:
        return zip_code or "—"
    if zip_code and zip_code not in address:
        return f"{address} ({zip_code})"
    return address


def build_case_address_summary(case: dict) -> str:
    addresses = _case_borrower_addresses(case)
    if not addresses:
        return "—"
    if len(addresses) == 1:
        return _format_address_with_zip(addresses[0])
    return "\n".join(f"{idx}. {_format_address_with_zip(entry)}" for idx, entry in enumerate(addresses, start=1))


def _latest_claim(case: dict) -> dict:
    payload = _parse_case_raw(case)
    detail = payload.get("detail") or {}
    claims = detail.get("claims") or payload.get("claims") or []
    if not claims:
        return {}
    return sorted(
        (claim for claim in claims if isinstance(claim, dict)),
        key=lambda claim: (
            str(claim.get("claim_date") or ""),
            str(claim.get("sent_at") or ""),
            str(claim.get("id") or ""),
        ),
        reverse=True,
    )[0]


def _normalize_address_token(token: str) -> str:
    text = " ".join((token or "").split()).strip(" ,")
    if not text:
        return ""
    for pattern, replacement in ADDRESS_ABBREVIATIONS:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    text = text.lower().title()
    return (
        text.replace("Г. ", "г. ")
        .replace("Д. ", "д. ")
        .replace("Аг. ", "аг. ")
        .replace("Пос. ", "пос. ")
        .replace("Ул. ", "ул. ")
        .replace("Пер. ", "пер. ")
        .replace("Пр-Т ", "пр-т ")
        .replace("Кв. ", "кв. ")
        .replace(" Область", " область")
        .replace(" Район", " район")
    )


def _split_address(address: str | None) -> dict[str, str | None]:
    parts = [
        _normalize_address_token(part)
        for part in str(address or "").split(",")
        if str(part).strip()
    ]
    city: str | None = None
    street_parts: list[str] = []
    extra: list[str] = []
    for part in parts:
        low = part.lower()
        if "беларус" in low:
            continue
        if city is None and low.startswith(LOCALITY_PREFIXES):
            city = part
            continue
        if low.startswith(("ул.", "пер.", "пр-т", "д.", "кв.")):
            street_parts.append(part)
        else:
            extra.append(part)
    street_line = ", ".join(street_parts) if street_parts else ", ".join(extra)
    return {"city": city, "street_line": street_line or None}


def _postal_lookup_meta(case: dict, target: dict[str, str] | None = None) -> dict[str, str | None]:
    payload = _parse_case_raw(case)
    lookup = payload.get("postal_lookup") or {}
    target_address = str((target or {}).get("address") or "").strip()
    primary_address = str(case.get("borrower_address") or "").strip()
    use_lookup = not target_address or target_address == primary_address
    postcode = str((target or {}).get("zip") or (lookup.get("postcode") if use_lookup else None) or case.get("borrower_zip") or "").strip() or None
    match_address = str(lookup.get("match_address") or "").strip()
    region = None
    locality = None
    if match_address and use_lookup:
        parts = [part.strip() for part in match_address.split(",") if part.strip()]
        if parts and re.fullmatch(r"\d{6}", parts[0]):
            parts = parts[1:]
        if parts:
            region = _normalize_address_token(parts[0])
        if len(parts) >= 2:
            locality = _normalize_address_token(parts[1])
    return {"postcode": postcode, "region": region, "locality": locality}


def _postal_address_lines(case: dict, target: dict[str, str]) -> list[str]:
    address_parts = _split_address(target.get("address"))
    lookup = _postal_lookup_meta(case, target)
    lines = [case.get("full_name") or "Получатель не указан"]
    if address_parts.get("street_line"):
        lines.append(str(address_parts["street_line"]))
    if address_parts.get("city") or lookup.get("locality"):
        lines.append(str(address_parts.get("city") or lookup.get("locality")))
    zip_region = ", ".join(part for part in [lookup.get("postcode"), lookup.get("region")] if part)
    if zip_region:
        lines.append(zip_region)
    if case.get("borrower_phone"):
        lines.append(f"Тел: {case['borrower_phone']}")
    return lines


def _debtor_header_lines(case: dict) -> list[str]:
    lines = [case.get("full_name") or "—"]
    addresses = _case_borrower_addresses(case)
    for idx, target in enumerate(addresses, start=1):
        prefix = "Адрес" if len(addresses) == 1 else f"Адрес {idx}"
        lines.append(f"{prefix}: {_format_address_with_zip(target)}")
    if case.get("borrower_phone"):
        lines.append(f"Тел: {case['borrower_phone']}")
    if case.get("borrower_email"):
        lines.append(f"Email: {case['borrower_email']}")
    return lines


def _build_contacts_block(name: str | None, address: str | None, phone: str | None, email: str | None) -> list[str]:
    lines = [name or "—"]
    if address:
        lines.append(address)
    if phone:
        lines.append(f"Тел: {phone}")
    if email:
        lines.append(f"Email: {email}")
    return lines


def collect_sms_missing_fields(case: dict, creditor: dict | None) -> list[str]:
    del creditor
    missing: list[str] = []
    if not case.get("full_name"):
        missing.append("ФИО заемщика")
    if not case.get("total_due"):
        missing.append("сумма долга")
    if not _loan_ref(case) or _loan_ref(case) == "—":
        missing.append("номер договора / займа")
    if not case.get("issued_at"):
        missing.append("дата договора / займа")
    return list(dict.fromkeys(missing))


def collect_claim_missing_fields(case: dict, creditor: dict | None, signature: dict | None) -> list[str]:
    missing = collect_sms_missing_fields(case, creditor)
    borrower_addresses = _case_borrower_addresses(case)
    if not creditor or not creditor.get("full_name"):
        missing.append("ФИО кредитора")
    if not case.get("document_id"):
        missing.append("ИН заемщика")
    if not borrower_addresses:
        missing.append("адрес заемщика")
    if not any(str(item.get("zip") or "").strip() for item in borrower_addresses):
        missing.append("ZIP-код заемщика")
    if not case.get("issued_at"):
        missing.append("дата договора / займа")
    if not creditor or not creditor.get("address"):
        missing.append("адрес кредитора")
    if not signature or not signature.get("file_path"):
        missing.append("подпись пользователя")
    return list(dict.fromkeys(missing))


def _join_non_empty(parts: list[str | None]) -> str:
    return ", ".join(part for part in parts if part)


def build_sms_text(case: dict, creditor: dict, variant: str = "soft") -> str:
    del creditor
    prefix = f"{_sms_name(case.get('full_name'))}, {_sms_service(case)} займ {_sms_ref(case)} от {_sms_date(case)}, долг {_sms_amount(case)}."
    soft_options = [
        " Оплатите добровольно, иначе последует обращение в суд с взысканием расходов.",
        " Оплатите добровольно, иначе обращусь в суд с взысканием расходов.",
        " Оплатите добровольно, иначе последует суд и взыскание расходов.",
    ]
    hard_options = [
        " При неоплате в течение недели обращаюсь в суд для взыскания. Последнее предупреждение.",
        " При неоплате за неделю обращаюсь в суд для взыскания. Последнее предупреждение.",
        " Неделя на оплату, далее обращаюсь в суд. Последнее предупреждение.",
        " Неделя на оплату, далее суд. Последнее предупреждение.",
    ]
    options = hard_options if variant == "hard" else soft_options
    for suffix in options:
        text = _fit_sms(prefix + suffix)
        if len(text) <= SMS_MAX_LEN:
            return text
    return _fit_sms(prefix)


def build_postal_address_text(case: dict) -> str:
    targets = _case_borrower_addresses(case)
    if not targets:
        targets = [{"address": str(case.get("borrower_address") or "").strip(), "zip": str(case.get("borrower_zip") or "").strip()}]
    header = "📮 <b>Адрес для отправки через Белпочту</b>" if len(targets) == 1 else "📮 <b>Адреса для отправки через Белпочту</b>"
    blocks = []
    for idx, target in enumerate(targets, start=1):
        prefix = [] if len(targets) == 1 else [f"[Адрес {idx}]"]
        blocks.append("\n".join(prefix + _postal_address_lines(case, target)))
    return header + "\n\n<pre>" + "\n\n".join(blocks) + "</pre>"


def _claim_amount_rows(case: dict) -> tuple[str, list[tuple[str, str, bool]]]:
    claim = _latest_claim(case)
    snapshot_date = _date(claim.get("claim_date") or datetime.now().isoformat())
    principal = _number(case.get("principal_outstanding"))
    if principal is None:
        principal = _number(case.get("amount")) or 0.0
    percent = _number(case.get("accrued_percent")) or 0.0
    fine = _number(case.get("fine_outstanding")) or 0.0
    total_value = _number(claim.get("amount"))
    if total_value is None:
        total_value = _number(case.get("total_due"))
    if total_value is None:
        total_value = principal + percent + fine
    return snapshot_date, [
        ("сумма невозвращенного в срок займа", _money(principal), False),
        ("проценты за пользование займом", _money(percent), False),
        ("пени", _money(fine), False),
        ("ИТОГО", _money(total_value), True),
    ]


def _claim_blocks(case: dict, creditor: dict) -> list[dict]:
    del creditor
    service_name = _sms_service(case)
    service_url = _service_url(case)
    deadline = _claim_deadline_date()
    snapshot_date, debt_rows = _claim_amount_rows(case)
    amount_total = debt_rows[-1][1]
    consequence_items = [
        "возложение судом на вас обязанности возместить мои судебные расходы по оплате государственной пошлины, адвокатской помощи и иные издержки, связанные с рассмотрением дела в суде;",
        "увеличение суммы долга по процентам за пользование займом и пеням, начисленным и подлежащим оплате за весь фактический период вашей просрочки возврата займа до дня возврата всей суммы займа включительно;",
        "взыскание с вас принудительного сбора в бюджет судебным исполнителем на стадии принудительного исполнения исполнительного документа, а также иных расходов по исполнению исполнительного документа;",
        "наложение ареста на ваше имущество, ограничение права на выезд из Республики Беларусь, ограничение права управления транспортным средством и иные меры, необходимые для обеспечения исполнения решения суда;",
        "обращение взыскания на ваше имущество, заработную плату и иные доходы, а также направление исполнительного документа по месту работы для удержаний;",
        "обращение взыскания на денежные средства на ваших банковских счетах и электронных кошельках, а также списание средств в счет погашения задолженности в рамках исполнительного производства;",
        "ухудшение вашей кредитной истории и, как следствие, сложности с получением в будущем кредита, займа или лизинга.",
    ]
    payment_sentence = (
        f"Оплатить задолженность возможно через сервис {service_url}."
        if service_url
        else "Оплатить задолженность необходимо в добровольном порядке в полном объеме."
    )
    return [
        {
            "type": "paragraph",
            "text": (
                f"Направляю настоящую претензию по договору займа № {_loan_ref(case)} от {_date(case.get('issued_at'))}, "
                f"оформленному через сервис {service_name}, с требованием добровольно погасить задолженность в досудебном порядке."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "**Настоящая претензия является последней возможностью урегулировать вопрос без подачи иска в суд.** "
                "После передачи материалов в суд сумма требований будет определяться уже на дату обращения в суд, "
                "а размер ваших расходов возрастет."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                f"По указанному договору вам была предоставлена сумма займа в размере {_money(case.get('amount'))}. "
                f"Срок возврата займа и оплаты начисленных процентов наступил {_date(case.get('due_at'))}, "
                "однако обязательства по договору вами не исполнены."
            ),
        },
        {
            "type": "paragraph",
            "text": f"Ваша задолженность по договору по состоянию на {snapshot_date} составляет:",
        },
        {"type": "debt_table", "rows": debt_rows},
        {
            "type": "paragraph",
            "text": (
                "На основании условий договора займа, а также статей 290, 311, 760 и 762 "
                "Гражданского кодекса Республики Беларусь требую добровольно погасить образовавшуюся задолженность."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                f"Требую не позднее {deadline} погасить задолженность по договору в общей сумме {amount_total}. "
                f"{payment_sentence}"
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "**Добровольное исполнение требований до обращения в суд является для вас финансово наиболее выгодным вариантом.** "
                "После начала судебного и исполнительного производства возможность урегулировать вопрос на более мягких условиях "
                "может быть утрачена."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "**В случае неисполнения** требований в добровольном порядке задолженность будет предъявлена "
                "к **судебному** и последующему **принудительному взысканию**."
            ),
        },
        {
            "type": "paragraph",
            "text": "**Обращаю ваше внимание:** обращение в суд и к принудительному взысканию повлечет для вас:",
        },
        {"type": "bullets", "items": consequence_items},
        {
            "type": "paragraph",
            "text": (
                f"В пределах срока, указанного в претензии (до {deadline}), готов рассмотреть ваши конструктивные "
                "предложения по добровольному урегулированию вопроса и полному погашению задолженности."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                f"По истечении указанного срока, то есть после {deadline}, при отсутствии оплаты либо приемлемых "
                "предложений задолженность будет взыскана в судебном и последующем принудительном порядке "
                "с увеличением суммы требований и взысканием дополнительных расходов."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                "**Дополнительно предупреждаю:** при дальнейшем уклонении от исполнения обязательств информация о вашей "
                "платежной дисциплине и наличии судебного взыскания будет объективно свидетельствовать о высокой степени "
                "риска при последующей оценке вашей кредитоспособности."
            ),
        },
        {
            "type": "paragraph",
            "text": (
                f"**ИТОГОВАЯ СУММА К ОПЛАТЕ НА {snapshot_date}: {amount_total}. "
                "ПРИ НЕОПЛАТЕ Я БУДУ ВЫНУЖДЕН НЕЗАМЕДЛИТЕЛЬНО ОБРАТИТЬСЯ В СУД ЗА ВЗЫСКАНИЕМ ДОЛГА, "
                "ГОСПОШЛИНЫ, РАСХОДОВ НА ПРЕДСТАВИТЕЛЯ И ИНЫХ ИЗДЕРЖЕК, А ТАКЖЕ ИНИЦИИРОВАТЬ "
                "ПОСЛЕДУЮЩЕЕ ПРИНУДИТЕЛЬНОЕ ИСПОЛНЕНИЕ.**"
            ),
        },
    ]


def build_claim_text(case: dict, creditor: dict) -> str:
    sections: list[str] = []
    for block in _claim_blocks(case, creditor):
        if block["type"] == "paragraph":
            sections.append(re.sub(r"\*\*(.*?)\*\*", r"\1", block["text"]).strip())
        elif block["type"] == "debt_table":
            sections.append("\n".join(f"{label}: {value}" for label, value, _ in block["rows"]))
        elif block["type"] == "bullets":
            sections.append("\n".join(f"- {item}" for item in block["items"]))
    return "\n\n".join(part for part in sections if part).strip()


def _claim_titles(case: dict) -> tuple[str, str]:
    del case
    return "ПРЕТЕНЗИЯ", "о добровольном урегулировании задолженности"


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(CLAIM_FONT_SIZE_PT)
    section = doc.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)


def _add_block(doc: Document, title: str, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(CLAIM_FONT_SIZE_PT)
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(CLAIM_FONT_SIZE_PT)


def _add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**") and len(part) >= 4
        content = part[2:-2] if is_bold else part
        run = p.add_run(content)
        run.font.size = Pt(CLAIM_FONT_SIZE_PT)
        run.bold = is_bold
    return p


def _add_bullet_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.first_line_indent = Mm(-4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    bullet = p.add_run("- ")
    bullet.font.size = Pt(CLAIM_FONT_SIZE_PT)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**") and len(part) >= 4
        content = part[2:-2] if is_bold else part
        run = p.add_run(content)
        run.font.size = Pt(CLAIM_FONT_SIZE_PT)
        run.bold = is_bold
    return p


def _add_debt_table(doc: Document, rows: list[tuple[str, str, bool]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for label, value, is_total in rows:
        cells = table.add_row().cells
        cells[0].width = Mm(115)
        cells[1].width = Mm(40)
        left_par = cells[0].paragraphs[0]
        right_par = cells[1].paragraphs[0]
        right_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        left_run = left_par.add_run(label)
        right_run = right_par.add_run(value)
        for run in (left_run, right_run):
            run.font.size = Pt(CLAIM_FONT_SIZE_PT)
            run.bold = is_total
    return table


def _cleanup_generated_docs() -> None:
    if not GENERATED_DOCS_DIR.exists():
        return
    cutoff = datetime.now() - timedelta(days=GENERATED_DOCS_TTL_DAYS)
    for path in GENERATED_DOCS_DIR.rglob("*"):
        if not path.is_file():
            continue
        modified = datetime.fromtimestamp(path.stat().st_mtime)
        if modified < cutoff:
            try:
                path.unlink()
            except OSError:
                pass
    for directory in sorted((path for path in GENERATED_DOCS_DIR.rglob("*") if path.is_dir()), reverse=True):
        try:
            next(directory.iterdir())
        except StopIteration:
            try:
                directory.rmdir()
            except OSError:
                pass


def render_claim_docx(case: dict, creditor: dict, signature_path: str) -> tuple[Path, str]:
    _cleanup_generated_docs()
    GENERATED_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    case_dir = GENERATED_DOCS_DIR / f"chat_{case['chat_id']}" / f"case_{case['id']}"
    case_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = case_dir / f"claim_{case['service']}_{case['id']}_{timestamp}.docx"

    doc = Document()
    _set_default_style(doc)

    debtor_lines = _debtor_header_lines(case)
    creditor_lines = _build_contacts_block(
        creditor.get("full_name"),
        creditor.get("address"),
        creditor.get("phone"),
        creditor.get("email"),
    )

    _add_block(doc, "Кому:", debtor_lines)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    _add_block(doc, "От:", creditor_lines)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    claim_title, claim_subtitle = _claim_titles(case)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(claim_title)
    run.bold = True
    run.font.size = Pt(15)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    run = subtitle.add_run(claim_subtitle)
    run.bold = True
    run.font.size = Pt(CLAIM_FONT_SIZE_PT)

    subject = doc.add_paragraph()
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject.paragraph_format.space_after = Pt(8)
    run = subject.add_run(f"по договору займа № {_loan_ref(case)} от {_date(case.get('issued_at'))}")
    run.bold = True
    run.font.size = Pt(CLAIM_FONT_SIZE_PT)

    body_paragraphs = []
    for block in _claim_blocks(case, creditor):
        if block["type"] == "paragraph":
            body_paragraphs.append(_add_body_paragraph(doc, block["text"]))
        elif block["type"] == "debt_table":
            _add_debt_table(doc, block["rows"])
        elif block["type"] == "bullets":
            for item in block["items"]:
                body_paragraphs.append(_add_bullet_paragraph(doc, item))

    for paragraph in body_paragraphs[-2:]:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.keep_with_next = True
    footer = doc.add_table(rows=1, cols=3)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.autofit = False
    left_cell, sign_cell, right_cell = footer.rows[0].cells
    left_cell.width = Mm(25)
    sign_cell.width = Mm(125)
    right_cell.width = Mm(25)
    for cell in (left_cell, sign_cell, right_cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left_par = left_cell.paragraphs[0]
    left_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_par.add_run(datetime.now().strftime("%d.%m.%Y"))
    left_run.font.size = Pt(CLAIM_FONT_SIZE_PT)

    sign_par = sign_cell.paragraphs[0]
    sign_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_run = sign_par.add_run()
    sign_run.add_picture(str(signature_path), width=Mm(25))

    right_par = right_cell.paragraphs[0]
    right_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_par.add_run(creditor.get("full_name") or "")
    right_run.font.size = Pt(CLAIM_FONT_SIZE_PT)

    doc.save(out_path)
    claim_text = build_claim_text(case, creditor)
    return out_path, claim_text


def serialize_case_payload(case: dict, creditor: dict | None = None) -> dict:
    return {
        "case": case,
        "creditor": creditor,
        "generated_at": datetime.now().isoformat(),
    }


def dump_payload(payload: dict) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str)
