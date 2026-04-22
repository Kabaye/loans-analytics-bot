from __future__ import annotations

from datetime import datetime
from pathlib import Path
import json
import logging

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from bot import config

log = logging.getLogger(__name__)

SERVICE_URLS = {
    "kapusta": "https://kapusta.by",
    "finkit": "https://finkit.by",
    "zaimis": "https://zaimis.by",
}

GENERATED_DOCS_DIR = Path(config.BASE_DIR) / "data" / "generated-docs"


def _money(value: float | int | None) -> str:
    try:
        return f"{float(value or 0):.2f} BYN"
    except (TypeError, ValueError):
        return "0.00 BYN"


def _parse_date(value: str | None) -> datetime | None:
    if not value:
        return None
    for candidate in (value, value.replace("Z", "+00:00")):
        try:
            return datetime.fromisoformat(candidate)
        except ValueError:
            continue
    return None


def _date(value: str | None) -> str:
    dt = _parse_date(value)
    if dt:
        return dt.strftime("%d.%m.%Y")
    return value or "—"


def _service_url(case: dict) -> str:
    return SERVICE_URLS.get(case.get("service"), "")


def _loan_ref(case: dict) -> str:
    return case.get("loan_number") or case.get("loan_id") or case.get("external_id") or "—"


def _voluntary_term(case: dict) -> str:
    days = case.get("voluntary_term_days")
    return str(days) if days else "—"


def _build_contacts_block(name: str | None, address: str | None, phone: str | None, email: str | None) -> list[str]:
    lines = [name or "—"]
    if address:
        lines.append(address)
    if phone:
        lines.append(f"Тел: {phone}")
    if email:
        lines.append(f"Email: {email}")
    return lines


def collect_sms_missing_fields(case: dict, creditor: dict | None) -> list[str]:
    missing: list[str] = []
    if not case.get("full_name"):
        missing.append("ФИО заемщика")
    if not case.get("total_due"):
        missing.append("сумма долга")
    if not _loan_ref(case) or _loan_ref(case) == "—":
        missing.append("номер договора / займа")
    if not case.get("issued_at"):
        missing.append("дата договора / займа")
    if not case.get("voluntary_term_days"):
        missing.append("срок добровольного погашения")
    if not creditor or not creditor.get("full_name"):
        missing.append("ФИО / название кредитора")
    if not creditor or (not creditor.get("phone") and not creditor.get("email")):
        missing.append("контакты кредитора")
    return list(dict.fromkeys(missing))


def collect_claim_missing_fields(case: dict, creditor: dict | None, signature: dict | None) -> list[str]:
    missing = collect_sms_missing_fields(case, creditor)
    if not case.get("document_id"):
        missing.append("ИН заемщика")
    if not case.get("borrower_address"):
        missing.append("адрес заемщика")
    if not case.get("borrower_zip"):
        missing.append("ZIP-код заемщика")
    if not case.get("issued_at"):
        missing.append("дата договора / займа")
    if not creditor or not creditor.get("address"):
        missing.append("адрес кредитора")
    if not signature or not signature.get("file_path"):
        missing.append("подпись пользователя")
    return list(dict.fromkeys(missing))


def build_sms_text(case: dict, creditor: dict) -> str:
    debt_breakdown = []
    if case.get("principal_outstanding") is not None:
        debt_breakdown.append(f"осн. долг {_money(case.get('principal_outstanding'))}")
    if case.get("accrued_percent") is not None:
        debt_breakdown.append(f"проценты {_money(case.get('accrued_percent'))}")
    if case.get("fine_outstanding") is not None:
        debt_breakdown.append(f"пеня {_money(case.get('fine_outstanding'))}")

    breakdown = ", ".join(debt_breakdown) if debt_breakdown else "структура долга уточняется"
    contacts = ", ".join(part for part in [creditor.get("phone"), creditor.get("email")] if part)

    return (
        f"Отправитель: {creditor.get('sms_sender') or creditor.get('full_name')}\n"
        f"Заемщик: {case.get('full_name')}\n"
        f"Сумма долга: {_money(case.get('total_due'))}\n"
        f"Договор / займ: {_loan_ref(case)} от {_date(case.get('issued_at'))}\n"
        f"Срок на добровольное погашение: {_voluntary_term(case)} дн.\n"
        f"Контакты для связи: {contacts or '—'}\n\n"
        f"У вас имеется задолженность по договору займа. Просим добровольно погасить долг "
        f"в течение {_voluntary_term(case)} дней. Текущий расчет: {breakdown}. "
        f"При непогашении задолженности будем вынуждены перейти к судебному взысканию."
    )


def build_claim_text(case: dict, creditor: dict) -> str:
    service_url = _service_url(case)
    return "\n\n".join([
        "ПРЕДЛОЖЕНИЕ\nо добровольном урегулировании задолженности",
        (
            f"Направляю настоящее письмо-предложение по договору займа {_loan_ref(case)} "
            f"от {_date(case.get('issued_at'))} через сервис онлайн-заимствования {service_url}, "
            "с целью урегулирования задолженности в добровольном досудебном порядке."
        ),
        (
            f"Убедительно прошу Вас погасить задолженность в полном объеме в течение "
            f"{_voluntary_term(case)} дней с даты получения настоящего документа."
        ),
        (
            "Обращаю Ваше внимание на то, что при передаче дела в суд размер требований "
            "будет увеличен за счет дальнейшего начисления процентов, пени и судебных расходов."
        ),
        (
            f"Сумма займа: {_money(case.get('amount'))} | "
            f"Проценты: {_money(case.get('accrued_percent'))} | "
            f"Пени: {_money(case.get('fine_outstanding'))}\n"
            f"ИТОГО к оплате: {_money(case.get('total_due'))} на дату отправки документа."
        ),
        (
            f"Для связи и добровольного урегулирования: {creditor.get('full_name')}, "
            f"{creditor.get('phone') or '—'}, {creditor.get('email') or '—'}."
        ),
    ])


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _add_block(doc: Document, title: str, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        run.font.size = Pt(12)


def _add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(10)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(12)


def render_claim_docx(case: dict, creditor: dict, signature_path: str) -> tuple[Path, str]:
    GENERATED_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    case_dir = GENERATED_DOCS_DIR / f"chat_{case['chat_id']}" / f"case_{case['id']}"
    case_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = case_dir / f"claim_{case['service']}_{case['id']}_{timestamp}.docx"

    doc = Document()
    _set_default_style(doc)

    debtor_lines = _build_contacts_block(
        case.get("full_name"),
        " ".join(part for part in [case.get("borrower_zip"), case.get("borrower_address")] if part),
        case.get("borrower_phone"),
        case.get("borrower_email"),
    )
    creditor_lines = _build_contacts_block(
        creditor.get("full_name"),
        creditor.get("address"),
        creditor.get("phone"),
        creditor.get("email"),
    )

    _add_block(doc, "Кому:", debtor_lines)
    doc.add_paragraph()
    _add_block(doc, "От:", creditor_lines)
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПРЕДЛОЖЕНИЕ")
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("о добровольном урегулировании задолженности")
    run.bold = True
    run.font.size = Pt(12)

    for paragraph in build_claim_text(case, creditor).split("\n\n"):
        _add_body_paragraph(doc, paragraph)

    totals = doc.add_table(rows=2, cols=2)
    totals.alignment = WD_TABLE_ALIGNMENT.CENTER
    totals.cell(0, 0).text = "Структура долга"
    totals.cell(0, 1).text = (
        f"Сумма займа: {_money(case.get('amount'))}\n"
        f"Проценты: {_money(case.get('accrued_percent'))}\n"
        f"Пени: {_money(case.get('fine_outstanding'))}"
    )
    totals.cell(1, 0).text = "Итого к оплате"
    totals.cell(1, 1).text = _money(case.get("total_due"))

    doc.add_paragraph()
    footer = doc.add_table(rows=1, cols=3)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.cell(0, 0).text = datetime.now().strftime("%d.%m.%Y")
    sign_par = footer.cell(0, 1).paragraphs[0]
    sign_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_run = sign_par.add_run()
    sign_run.add_picture(str(signature_path), width=Mm(35))
    footer.cell(0, 2).text = creditor.get("full_name") or ""

    doc.save(out_path)
    claim_text = build_claim_text(case, creditor)
    return out_path, claim_text


def serialize_case_payload(case: dict, creditor: dict | None = None) -> dict:
    return {
        "case": case,
        "creditor": creditor,
        "generated_at": datetime.now().isoformat(),
    }


def dump_payload(payload: dict) -> str:
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str)
